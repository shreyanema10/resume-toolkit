#!/usr/bin/env python3
"""
Resume generator for Shreya Nema.

Usage:
  python3 generate_resume.py                        # base resume (no company)
  python3 generate_resume.py "Google"               # company-specific output
  python3 generate_resume.py "Google" --pdf-only
  python3 generate_resume.py "Google" --docx-only

Reads:  resume_data.json  (in same directory)
Output: Company Resume/<Company>/<filename>.pdf + .docx
        If no company, outputs to LeaveHiLabs/ directory as base files.
"""

import json
import sys
import os
import re
from pathlib import Path
from datetime import datetime

SCRIPT_DIR = Path(__file__).parent
DATA_FILE  = SCRIPT_DIR / "resume_data.json"

BLUE_HEX   = "1F5887"
BLUE_RGB   = (31, 88, 135)
TABLE_BG   = "BDD7EE"

MONTHS = {
    1:"Jan", 2:"Feb", 3:"Mar", 4:"Apr", 5:"May", 6:"Jun",
    7:"Jul", 8:"Aug", 9:"Sep", 10:"Oct", 11:"Nov", 12:"Dec"
}


# ── Path helpers ──────────────────────────────────────────────────────────────

def output_paths(company: str | None):
    month = MONTHS[datetime.now().month]
    if company:
        safe = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
        folder = SCRIPT_DIR / "Company Resume" / safe
        folder.mkdir(parents=True, exist_ok=True)
        stem = f"Resume_Shreya_Nema_{month}_2026_{safe}"
    else:
        folder = SCRIPT_DIR
        stem = "Shreya_Nema_Resume_AI_PM"
    return folder / f"{stem}.docx", folder / f"{stem}.pdf"


# ── Data ──────────────────────────────────────────────────────────────────────

def load_resume():
    with open(DATA_FILE, encoding="utf-8") as f:
        return json.load(f)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def generate_docx(resume: dict, out_path: Path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE = RGBColor(*BLUE_RGB)

    doc = Document()

    # Global normal style baseline
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    # Page margins
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin   = Inches(0.6)
    sec.right_margin  = Inches(0.6)

    # EMU → twips: 1 inch = 914400 EMU = 1440 twips, so 1 EMU = 1440/914400 twips
    cw_emu   = sec.page_width - sec.left_margin - sec.right_margin
    cw_twips = int(cw_emu * 1440 / 914400)

    # ── XML helpers ────────────────────────────────────────────────────────────

    def _spacing(para, before_pt=0, after_pt=0):
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:spacing')):
            pPr.remove(old)
        el = OxmlElement('w:spacing')
        el.set(qn('w:before'), str(int(before_pt * 20)))
        el.set(qn('w:after'),  str(int(after_pt  * 20)))
        pPr.append(el)

    def _bottom_border(para):
        pPr = para._p.get_or_add_pPr()
        bdr = OxmlElement('w:pBdr')
        bt  = OxmlElement('w:bottom')
        bt.set(qn('w:val'),   'single')
        bt.set(qn('w:sz'),    '6')
        bt.set(qn('w:space'), '1')
        bt.set(qn('w:color'), BLUE_HEX)
        bdr.append(bt)
        pPr.append(bdr)

    def _right_tab(para):
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab  = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), str(cw_twips))
        tabs.append(tab)
        pPr.append(tabs)

    def _hanging_indent(para, left_twips=230, hang_twips=180):
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:ind')):
            pPr.remove(old)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'),    str(left_twips))
        ind.set(qn('w:hanging'), str(hang_twips))
        pPr.append(ind)

    def _cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')):
            tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    # ── Run factory ────────────────────────────────────────────────────────────

    def run(para, text, size=9.5, bold=False, color=None):
        r = para.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.bold = bold
        if color:
            r.font.color.rgb = color
        return r

    # ── Section header ─────────────────────────────────────────────────────────

    def section_hdr(title, before_pt=5):
        p = doc.add_paragraph()
        _spacing(p, before_pt=before_pt, after_pt=2)
        run(p, title, size=10.5, bold=True, color=BLUE)
        _bottom_border(p)

    # ── Header ────────────────────────────────────────────────────────────────

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before_pt=0, after_pt=1)
    run(p, resume["name"], size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before_pt=0, after_pt=1)
    run(p, resume["title"], size=10.5, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before_pt=0, after_pt=4)
    run(p, resume["contact"], size=9.5)

    # ── Summary ───────────────────────────────────────────────────────────────

    section_hdr("SUMMARY", before_pt=3)
    p = doc.add_paragraph()
    _spacing(p, before_pt=1, after_pt=2)
    run(p, resume["summary"], size=9.5)

    # ── Experience ────────────────────────────────────────────────────────────

    section_hdr("PROFESSIONAL EXPERIENCE", before_pt=3)

    for i, job in enumerate(resume["experience"]):
        # Company + date line
        p = doc.add_paragraph()
        _spacing(p, before_pt=(3 if i > 0 else 1), after_pt=0)
        _right_tab(p)
        run(p, f"{job['company']}  ·  {job['location']}", size=10, bold=True)
        run(p, f"\t{job['start']} – {job['end']}", size=10)

        # Job title
        p = doc.add_paragraph()
        _spacing(p, before_pt=0, after_pt=1)
        run(p, job["title"], size=10, bold=True, color=BLUE)

        # Bullets
        for bullet in job["bullets"]:
            p = doc.add_paragraph()
            _spacing(p, before_pt=0, after_pt=1.5)
            _hanging_indent(p)
            run(p, "• ", size=9.5)
            run(p, bullet["bold"], size=9.5, bold=True)
            run(p, bullet["text"], size=9.5)

    # ── Technical Proficiency ─────────────────────────────────────────────────

    section_hdr("TECHNICAL PROFICIENCY", before_pt=3)

    tbl = doc.add_table(rows=len(resume["skills"]), cols=2)
    tbl.style = 'Table Grid'

    COL1_TWIPS = int(Inches(0.95).twips)
    COL2_TWIPS = cw_twips - COL1_TWIPS

    for i, skill in enumerate(resume["skills"]):
        r0 = tbl.rows[i]

        c0 = r0.cells[0]
        c0.paragraphs[0].clear()
        _spacing(c0.paragraphs[0], before_pt=1.5, after_pt=1.5)
        run(c0.paragraphs[0], skill["category"], size=9.5, bold=True, color=BLUE)
        _cell_bg(c0, TABLE_BG)

        c1 = r0.cells[1]
        c1.paragraphs[0].clear()
        _spacing(c1.paragraphs[0], before_pt=1.5, after_pt=1.5)
        run(c1.paragraphs[0], skill["content"], size=9.5)

        # Enforce widths via XML
        for cell, w in ((c0, COL1_TWIPS), (c1, COL2_TWIPS)):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW  = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # ── Education ─────────────────────────────────────────────────────────────

    section_hdr("EDUCATION", before_pt=3)

    edu = resume["education"]
    p = doc.add_paragraph()
    _spacing(p, before_pt=1, after_pt=0)
    _right_tab(p)
    run(p, edu["school"], size=10, bold=True)
    run(p, f"  |  {edu['dates']}", size=10)

    p = doc.add_paragraph()
    _spacing(p, before_pt=0, after_pt=0)
    run(p, edu["degree"], size=9.5)

    p = doc.add_paragraph()
    _spacing(p, before_pt=0, after_pt=2)
    r = run(p, edu["courses"], size=9)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # ── Leadership ────────────────────────────────────────────────────────────

    section_hdr("LEADERSHIP", before_pt=3)

    p = doc.add_paragraph()
    _spacing(p, before_pt=1, after_pt=0)
    _hanging_indent(p)
    run(p, "• ", size=9.5)
    run(p, resume["leadership"], size=9.5)

    doc.save(str(out_path))
    print(f"  DOCX → {out_path}")


# ── PDF ───────────────────────────────────────────────────────────────────────

def generate_pdf(resume: dict, out_path: Path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.colors import HexColor, black
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, HRFlowable, KeepTogether
    )
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT

    BLUE     = HexColor(f"#{BLUE_HEX}")
    TBL_BG   = HexColor(f"#{TABLE_BG}")
    GRAY     = HexColor("#595959")
    GRID     = HexColor("#AAAAAA")
    W, H     = letter
    LM = RM  = 0.55 * inch
    TM = BM  = 0.45 * inch
    CW       = W - LM - RM

    doc = SimpleDocTemplate(
        str(out_path), pagesize=letter,
        leftMargin=LM, rightMargin=RM,
        topMargin=TM,  bottomMargin=BM,
    )

    F    = 'Helvetica'
    FB   = 'Helvetica-Bold'

    def S(name, **kw):
        kw.setdefault('fontName', F)
        return ParagraphStyle(name, **kw)

    s_name     = S('name',    fontName=FB, fontSize=19,   alignment=TA_CENTER, spaceAfter=1,  spaceBefore=0, leading=22)
    s_title    = S('title',   fontSize=10,  textColor=BLUE, alignment=TA_CENTER, spaceAfter=1,  spaceBefore=0, leading=12)
    s_contact  = S('contact', fontSize=9.5, alignment=TA_CENTER, spaceAfter=4,  spaceBefore=0, leading=11)
    s_sec      = S('sec',     fontName=FB, fontSize=10,   textColor=BLUE, spaceBefore=4, spaceAfter=0, leading=12)
    s_summary  = S('summary', fontSize=9.5, alignment=TA_JUSTIFY, spaceAfter=2, spaceBefore=0, leading=11.5)
    s_co       = S('co',      fontName=FB, fontSize=9.5,  spaceAfter=0, spaceBefore=1, leading=11)
    s_date     = S('date',    fontSize=9.5, alignment=TA_RIGHT, spaceAfter=0, spaceBefore=1, leading=11)
    s_jobtitle = S('jt',      fontName=FB, fontSize=9.5,  textColor=BLUE, spaceAfter=0.5, spaceBefore=0, leading=11)
    s_bullet   = S('bullet',  fontSize=9.5, leftIndent=10, firstLineIndent=0, spaceAfter=1,   spaceBefore=0, leading=11.5)
    s_edu_hdr  = S('eduhdr',  fontName=FB, fontSize=9.5,  spaceAfter=0, spaceBefore=0, leading=11)
    s_edu_dt   = S('edudt',   fontSize=9.5, alignment=TA_RIGHT, spaceAfter=0, spaceBefore=0, leading=11)
    s_edu_deg  = S('edudeg',  fontSize=9.5, spaceAfter=0, spaceBefore=0, leading=11)
    s_courses  = S('courses', fontSize=9,   textColor=GRAY, spaceAfter=2, spaceBefore=0, leading=10.5)
    s_ch       = S('ch',      fontName=FB, fontSize=9.5, textColor=BLUE, spaceAfter=0, spaceBefore=0, leading=11)
    s_cb       = S('cb',      fontSize=9.5, spaceAfter=0, spaceBefore=0, leading=11)

    def hr():
        return HRFlowable(width='100%', thickness=0.5, color=BLUE, spaceAfter=1.5, spaceBefore=0)

    def sec_hdr(title):
        return [Paragraph(title, s_sec), hr()]

    def co_row(job):
        data = [[
            Paragraph(f"<b>{job['company']}  ·  {job['location']}</b>", s_co),
            Paragraph(f"{job['start']} – {job['end']}", s_date),
        ]]
        t = Table(data, colWidths=[CW * 0.68, CW * 0.32])
        t.setStyle(TableStyle([
            ('ALIGN',         (1, 0), (1, 0), 'RIGHT'),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING',   (0, 0), (-1, -1), 0),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
            ('TOPPADDING',    (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        return t

    def bullet_p(bold_text, rest_text):
        return Paragraph(f"•  <b>{bold_text}</b>{rest_text}", s_bullet)

    story = []

    # Header
    story.append(Paragraph(resume["name"], s_name))
    story.append(Paragraph(resume["title"], s_title))
    story.append(Paragraph(resume["contact"], s_contact))

    # Summary
    story += sec_hdr("SUMMARY")
    story.append(Paragraph(resume["summary"], s_summary))

    # Experience
    story += sec_hdr("PROFESSIONAL EXPERIENCE")

    for i, job in enumerate(resume["experience"]):
        block = []
        if i > 0:
            block.append(Spacer(1, 2))
        block.append(co_row(job))
        block.append(Paragraph(job["title"], s_jobtitle))
        for b in job["bullets"]:
            block.append(bullet_p(b["bold"], b["text"]))
        story.append(KeepTogether(block))

    # Technical Proficiency
    story += sec_hdr("TECHNICAL PROFICIENCY")

    COL1 = 0.95 * inch
    COL2 = CW - COL1
    rows = [[Paragraph(s["category"], s_ch), Paragraph(s["content"], s_cb)]
            for s in resume["skills"]]
    tbl = Table(rows, colWidths=[COL1, COL2])
    tbl.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (0, -1), TBL_BG),
        ('BOX',           (0, 0), (-1, -1), 0.5, GRID),
        ('INNERGRID',     (0, 0), (-1, -1), 0.5, GRID),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING',    (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING',   (0, 0), (-1, -1), 4),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 4),
    ]))
    story.append(tbl)

    # Education
    story += sec_hdr("EDUCATION")

    edu = resume["education"]
    edu_row = [[
        Paragraph(f"<b>{edu['school']}</b>", s_edu_hdr),
        Paragraph(f"| {edu['dates']}", s_edu_dt),
    ]]
    edu_t = Table(edu_row, colWidths=[CW * 0.70, CW * 0.30])
    edu_t.setStyle(TableStyle([
        ('ALIGN',         (1, 0), (1, 0), 'RIGHT'),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('TOPPADDING',    (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(edu_t)
    story.append(Paragraph(edu["degree"], s_edu_deg))
    story.append(Paragraph(edu["courses"], s_courses))

    # Leadership
    story += sec_hdr("LEADERSHIP")
    story.append(Paragraph(f"•  {resume['leadership']}", s_bullet))

    doc.build(story)
    print(f"  PDF  → {out_path}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    args = sys.argv[1:]

    # Filter flags
    pdf_only  = '--pdf-only'  in args
    docx_only = '--docx-only' in args
    flags     = {'--pdf-only', '--docx-only'}
    pos_args  = [a for a in args if a not in flags]

    company   = pos_args[0] if pos_args else None
    do_docx   = not pdf_only
    do_pdf    = not docx_only

    docx_path, pdf_path = output_paths(company)
    resume = load_resume()

    label = f"[{company}]" if company else "[base]"
    print(f"\nGenerating resume {label} ...")

    if do_docx:
        generate_docx(resume, docx_path)
    if do_pdf:
        generate_pdf(resume, pdf_path)

    print("Done.\n")


if __name__ == '__main__':
    main()
